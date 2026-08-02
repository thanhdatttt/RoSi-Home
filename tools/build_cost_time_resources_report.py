from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"D:\agy-cli-projects\PM-Group\RoSi-Home\docs\reports\RosiHome_Cost_Time_Resources_Report.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4CE"
PALE_RED = "FDE9E7"
GRAY = "666666"
WHITE = "FFFFFF"
GREEN = "2E7D32"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

def set_font(run, name="Calibri", size=11, bold=None, italic=None, color="000000"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    st = styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "Report Kicker" not in [s.name for s in styles]:
    kicker_style = styles.add_style("Report Kicker", WD_STYLE_TYPE.PARAGRAPH)
else:
    kicker_style = styles["Report Kicker"]
kicker_style.paragraph_format.space_after = Pt(4)
kicker_style.font.name = "Calibri"
kicker_style.font.size = Pt(10)
kicker_style.font.bold = True
kicker_style.font.color.rgb = RGBColor.from_string(BLUE)

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

def set_table_geometry(table, widths_dxa, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa)-1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

def add_table(headers, rows, widths, font_size=9.1, header_fill=LIGHT, aligns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(str(text)), size=font_size, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if aligns and i < len(aligns):
                p.alignment = aligns[i]
            set_font(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])

def setup_header_footer(section):
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_font(hp.add_run("RosiHome | Software Project Estimation & Control"), size=8.5, bold=True, color=GRAY)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    set_font(fp.add_run("Internal baseline | 22/07/2026 | Page "), size=8, color=GRAY)
    add_field(fp, "PAGE")

setup_header_footer(sec)

def add_title(text, size=25, color=NAVY, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size=size, bold=True, color=color)
    return p

def add_para(text="", bold_prefix=None, italic=False, color="000000", after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True, color=color)
        set_font(p.add_run(text[len(bold_prefix):]), italic=italic, color=color)
    else:
        set_font(p.add_run(text), italic=italic, color=color)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text))
    return p

def add_number(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text))
    return p

def callout(label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0,0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(label + " "), size=10.5, bold=True, color=NAVY)
    set_font(p.add_run(text), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

def heading(text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")

def page_break():
    doc.add_page_break()

# Cover / executive page
doc.add_paragraph("PROJECT MANAGEMENT BASELINE", style="Report Kicker")
add_title("RosiHome Cost, Time & Resources Report", 27)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
set_font(p.add_run("Actual AI-assisted delivery evidence, forecast model, planning baseline and monitoring controls"), size=13, italic=True, color=GRAY)

add_table(
    ["Document", "Value"],
    [
        ["Project", "RosiHome Property Management System"],
        ["Evidence cutoff", "22/07/2026 (Asia/Saigon)"],
        ["Scope basis", "Phân công.txt (BE/FE batch table) + Product Backlog"],
        ["Purpose", "Oral defense: estimation, planning, monitoring & control"],
        ["Status", "Working baseline - requires team validation before SOW/contract use"],
    ],
    [1900, 7460], font_size=9.6,
    aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT]
)

heading("Executive summary", 1)
callout("Measured evidence.", "Three backend work packages report 47.38 hours and 213.36 million tokens. Incremental AI cash outlay during the reported period is 0 VND because all named access was trial/free; this does not mean economic cost is zero.")
add_table(
    ["Metric", "Observed", "Interpretation"],
    [
        ["Recorded work time", "47.38 h", "Mix of telemetry agent time and member-reported elapsed time"],
        ["Reported tokens", "213.36 M", "Not normalized across tools; includes different token accounting methods"],
        ["Provisionally covered stories", "34 / 55 BE stories", "Only if reported batches are complete; not proof of global DoD"],
        ["Observed AI cash cost", "0 VND incremental", "Trial/free access; post-trial pricing remains an open input"],
        ["Proposal baseline", "8–10 weeks; 4.25 M VND", "Project-level schedule and tools/hosting budget, not yet reconciled to actuals"],
    ],
    [2350, 1900, 5110], font_size=9.0,
    aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
)
add_para("Management conclusion: the evidence supports an AI-accelerated backend build, but it does not yet support claiming the MVP is Done. Tests, review, integration, deployment and mobile acceptance must be reported separately.", bold_prefix="Management conclusion:")

page_break()
heading("1. Scope and batch baseline", 1)
add_para("The operative allocation is the BE/FE table in Phân công.txt. The repository assignments file is not used as the governing source for this report. Product meaning and acceptance criteria come from docs/product_backlog.md.")

add_table(
    ["Batch", "BE1 Chí", "BE2 Đạt", "BE3 Minh", "FE1 MXH", "FE2 Quân"],
    [
        ["1 Foundation", "Auth 01–06; Profile 01", "Property 01–02; Room 01–03", "Utility 01–02; Charge 01", "Auth/Profile/Property/Room UI", "Design system, navigation, shared components, Utility UI"],
        ["2 Core", "Tenant 01–02; Lease 01–06", "Meter 01–03", "Maintenance 01–05", "Tenant & Lease UI", "Meter & Maintenance UI"],
        ["3 Billing", "Review, test, bug fix", "Invoice 01–04", "VietQR 01–02; Payment 01–03; Reminder 01–02", "Invoice & Payment UI", "VietQR, proof upload, notification UI"],
        ["4 Analytics", "Dashboard 01–02", "Dashboard 03–04", "Report 01–05", "Dashboard UI", "Report UI"],
    ],
    [950, 1540, 1420, 1740, 1700, 2010], font_size=7.8,
    aligns=[WD_ALIGN_PARAGRAPH.CENTER]*6
)

heading("Dependency sequence", 2)
add_number("Batch 1 establishes authentication, property/room and billing/payment foundations.")
add_number("Batch 2 establishes tenant/lease, meter and maintenance core flows.")
add_number("Batch 3 consumes Batch 2 data for invoice, payment, VietQR and reminders.")
add_number("Batch 4 consumes system-wide data for dashboards and reports.")
callout("Planning rule.", "Parallelize work inside a batch only after API contracts are agreed. Move to the next batch through an explicit integration gate, not merely when coding stops.", PALE_GOLD)

heading("Definition of Done used by this report", 2)
add_bullet("All acceptance criteria pass on the Mobile delivery surface.")
add_bullet("Backend authorization and ownership rules are enforced.")
add_bullet("Automated success-path and critical validation/authorization tests pass.")
add_bullet("No unresolved critical/high defect remains; code is reviewed and merged.")
add_bullet("Migrations/configuration are reproducible; behavior is deployed and verified in the agreed environment.")

page_break()
heading("2. Data sources and measurement method", 1)
add_table(
    ["Source", "Evidence", "Confidence", "Limitation"],
    [
        ["Minh / BE3", "Codex telemetry by turn for US-MAINT-01→05", "High for time/token", "MAINT-03 was retrospectively split from MAINT-02; token total includes cached input"],
        ["Đạt / BE2 (inferred)", "3 h/6.3 M Batch 1; 5 h/13 M Batch 2+3; 2 h/2.5 M Batch 4", "Medium", "Member identity and story completion inferred from allocation; no raw logs"],
        ["Chí / BE1 (inferred)", "15 h Batch 1; 20 h Batch 2 + tests/FE setup/fixes/docs; 135 M tokens", "Medium-low", "Mixed technical tasks and stories; token not split by batch"],
        ["Proposal", "8–10 weeks; 4.25 M VND", "Planning baseline", "Contains scope statements that conflict with current backlog"],
        ["Product backlog", "Story descriptions, dependencies, AC and global DoD", "High for scope", "Stories are Refined; Refined is not automatically Ready or Done"],
    ],
    [1550, 3170, 1320, 3320], font_size=8.2
)

heading("Metric definitions", 2)
add_bullet("Agent coding time: active/elapsed time reported by telemetry or a member. It is not automatically human hands-on time.")
add_bullet("Wall-clock lead time: start-to-finish calendar duration including waiting, review and dependencies. It is not yet consistently recorded.")
add_bullet("Total tokens: provider/tool-reported token volume. It is a resource indicator, not a universal billing amount.")
add_bullet("Throughput: user stories reaching the agreed completion state per unit time. Coding complete and Done must be reported separately.")
add_bullet("Cash cost: actual money paid. Shadow labor cost: hours multiplied by an agreed hourly rate for economic comparison.")

callout("Comparability warning.", "Minh's telemetry includes cached-input accounting; the other members supplied aggregate totals without input/output/cache breakdown. Therefore, token-per-hour differences cannot be interpreted as model efficiency.", PALE_RED)

heading("Recommended collection record per story", 2)
add_table(
    ["Field", "Required value"],
    [
        ["Identity", "US-ID, owner, batch, BE/FE, Codex task/conversation ID"],
        ["Time", "Start/end, agent elapsed, human review time, blocked time, deploy time"],
        ["AI usage", "Model, input, cached input, output, reasoning, total token, subscription/API mode"],
        ["Quality", "Tests passed/failed, review result, defects, acceptance criteria pass count"],
        ["Delivery", "Commit, PR, merge, deployment URL/environment, verification timestamp"],
    ],
    [1900, 7460], font_size=9.1
)

page_break()
heading("3. Actual cost-time-resource evidence", 1)

hours_minh = 8569.2 / 3600
hours_dat = 10.0
hours_chi = 35.0
tokens_minh = 56.564626
tokens_dat = 21.8
tokens_chi = 135.0
total_hours = hours_minh + hours_dat + hours_chi
total_tokens = tokens_minh + tokens_dat + tokens_chi

add_table(
    ["Work package", "Scope represented", "Time", "Tokens", "Rate indicators"],
    [
        ["Minh / BE3", "US-MAINT-01→05 (5 stories)", f"{hours_minh:.2f} h", f"{tokens_minh:.2f} M", f"0.48 h/US; 11.31 M/US"],
        ["Đạt / BE2*", "Batch 1–4 allocation (14 stories)", "10.00 h", "21.80 M", "0.71 h/US; 1.56 M/US"],
        ["Chí / BE1*", "Infra + 3 named Auth stories + Batch 2 + tests/FE/fixes/docs", "35.00 h", "135.00 M", "3.86 M token/h; per-US rate not clean"],
        ["TOTAL", "Three reported backend packages", f"{total_hours:.2f} h", f"{total_tokens:.2f} M", "Observed aggregate; not normalized"],
    ],
    [1500, 3440, 1050, 1200, 2170], font_size=8.6,
    aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
)
add_para("* Member-to-column mapping is inferred from the described batch work and must be confirmed before the report becomes contractual.", italic=True, color=GRAY, after=8)

heading("Maintenance detail - directly measured sample", 2)
maint = [
    ("US-MAINT-01", "Submit a maintenance request", 17.390002, "38m 57s", "Measured"),
    ("US-MAINT-02", "View submitted maintenance requests", 5.765790, "25m 31s", "Estimated split"),
    ("US-MAINT-03", "Review maintenance requests", 8.435546, "21m 43s", "Estimated split"),
    ("US-MAINT-04", "Update maintenance status", 11.730275, "35m 17s", "Measured"),
    ("US-MAINT-05", "View maintenance history by room", 13.243013, "21m 23s", "Measured"),
]
add_table(
    ["Story", "Backlog outcome", "Tokens", "Agent time", "Evidence"],
    [[a,b,f"{c:.3f} M",d,e] for a,b,c,d,e in maint] + [["TOTAL","5 maintenance stories",f"{tokens_minh:.3f} M","2h 22m 50s","Telemetry total"]],
    [1450, 3450, 1300, 1350, 1810], font_size=8.7,
    aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
)

heading("Đạt / BE2 batch detail - member-reported sample", 2)
add_table(
    ["Batch", "Assigned stories", "Time", "Tokens", "Derived"],
    [
        ["1", "Property 01–02; Room 01–03 (5)", "3 h", "6.3 M", "0.60 h/US; 1.26 M/US"],
        ["2 + 3", "Meter 01–03; Invoice 01–04 (7)", "5 h", "13.0 M", "0.71 h/US; 1.86 M/US"],
        ["4", "Dashboard 03–04 (2)", "2 h", "2.5 M", "1.00 h/US; 1.25 M/US"],
        ["TOTAL", "14 assigned stories", "10 h", "21.8 M", "0.71 h/US; 1.56 M/US"],
    ],
    [1000, 3500, 1050, 1200, 2610], font_size=8.8
)

page_break()
heading("4. Interpretation and productivity analysis", 1)
heading("What the measurements do support", 2)
add_bullet("AI agents can complete bounded backend story packages in hours rather than the proposal's multi-week coding phases.")
add_bullet("Batching by dependency domain allows parallel work: Auth, Property/Room and Billing Foundation can proceed independently in Batch 1.")
add_bullet("The largest time package includes infrastructure, tests, frontend setup, bug fixing and documentation, demonstrating that non-story work materially affects delivery effort.")

heading("What the measurements do not support", 2)
add_bullet("They do not prove all acceptance criteria passed or that every story was deployed successfully.")
add_bullet("They do not prove one model is more efficient: token definitions and task complexity differ.")
add_bullet("They do not measure frontend, integration, stakeholder feedback, production operations or blocked time consistently.")
add_bullet("They do not justify multiplying a single average across the whole backlog without complexity and dependency buffers.")

heading("Reference rates", 2)
add_table(
    ["Rate", "Value", "Use", "Caution"],
    [
        ["Story-only weighted rate", "0.65 h/US", "Optimistic backend coding benchmark from Minh + Đạt", "Excludes Chí's technical overhead and full DoD"],
        ["Team blended rate", "1.39 h/claimed US", "Most-likely planning anchor for 34 provisional stories", "Depends on provisional story count"],
        ["Overhead-heavy rate", "2.33 h/US-equivalent", "Pessimistic bound from Chí's mixed package", "Not a pure per-story measure"],
        ["Token central reference", "4.12 M/US", "Minh + Đạt story-only aggregate", "Very wide observed range; not a cost quote"],
    ],
    [2100, 1800, 2870, 2590], font_size=8.8
)

heading("Three-point estimate", 2)
add_para("For planning only, use O = 0.65 h/US, M = 1.39 h/US and P = 2.33 h/US. PERT expected effort is (O + 4M + P) / 6 = 1.42 h per remaining backend story before integration and contingency.")
callout("Recommended estimator.", "Use user-story throughput for schedule, token volume for AI capacity/cost monitoring, and quality/deployment evidence for completion. Do not use token count alone as a proxy for progress.")

page_break()
heading("5. Forward estimate and resource plan", 1)
add_para("A provisional 34 of 55 backend stories are represented by the supplied work claims: Minh 5 Maintenance; Đạt 14 assigned stories; Chí 3 named Auth stories plus 12 Batch-2 Tenant/Lease stories. This leaves 21 nominal stories. The count must be reconciled against PR, test and deployment evidence.")

remaining = 21
base = remaining * 1.42
integration = base * 1.30
contingency = integration * 1.20
add_table(
    ["Forecast layer", "Formula", "Effort", "Meaning"],
    [
        ["Base coding", "21 × 1.42 h", f"{base:.1f} h", "Expected agent-supported implementation"],
        ["Integration & DoD", "Base × 1.30", f"{integration:.1f} h", "Review, tests, migrations, integration, deployment verification"],
        ["Planning commitment", "Integrated × 1.20", f"{contingency:.1f} h", "Risk reserve for ambiguity, defects and rework"],
        ["Ideal 3-BE parallel lead time", "Commitment ÷ 3", f"{contingency/3:.1f} h", "Theoretical; dependency and review queues increase calendar time"],
    ],
    [1800, 2000, 1350, 4210], font_size=8.8
)
add_para("Recommended calendar commitment for remaining backend scope: 3–5 focused working days with three backend members, provided API contracts are frozen and no high-severity integration defect appears. This forecast excludes unmeasured frontend completion, pilot feedback and long-running operational support.", bold_prefix="Recommended calendar commitment for remaining backend scope:")

heading("Token capacity forecast", 2)
add_table(
    ["Scenario", "Rate", "21-story volume", "Use"],
    [
        ["Low", "1.56 M/US", "≈ 32.8 M", "Compact contexts and CRUD-like stories"],
        ["Central", "4.12 M/US", "≈ 86.6 M", "Story-only observed aggregate"],
        ["High", "11.31 M/US", "≈ 237.6 M", "Large cached contexts, repeated review/debug cycles"],
    ],
    [1550, 1700, 2100, 4010], font_size=9.0
)
add_para("The range is intentionally wide. It should be narrowed after another 5–10 stories are measured using one telemetry schema and the same definition of total token.", italic=True, color=GRAY)

heading("Resource allocation", 2)
add_table(
    ["Resource", "Current role", "Planning control"],
    [
        ["BE1 Chí", "Auth; Tenant/Lease; Dashboard 01–02; integration support", "Limit WIP to one story plus one review item"],
        ["BE2 Đạt", "Property/Room; Meter; Invoice; Dashboard 03–04", "Own Batch 3 invoice integration gate"],
        ["BE3 Minh", "Utility/Charge; Maintenance; Payment/Reminder; Report", "Own telemetry and report data quality"],
        ["FE1/FE2", "Mobile UI and integration", "Record separate FE story time; do not merge with BE throughput"],
        ["AI assistants", "Generate/refine code, tests, docs and diagnostics", "One story per task; human review required"],
        ["Platform", "GitHub, PostgreSQL, Docker/cloud, CI, telemetry", "Track uptime, deploy cost and service consumption"],
    ],
    [1750, 3900, 3710], font_size=8.7
)

page_break()
heading("6. Cost model", 1)
heading("Observed cash cost", 2)
add_para("For the measured work, the incremental AI assistant cash cost is recorded as 0 VND: GPT-5.6 Sol was accessed through a one-month Plus trial and Hy3 was reported as free. Subscription entitlement has no verified post-trial price in the supplied evidence.")
callout("Cost statement.", "Free trial means zero current cash outflow, not zero resource consumption. Tokens, developer time, compute, hosting and future subscription renewal remain economically relevant.", PALE_GOLD)

heading("Shadow labor scenarios", 2)
rows = []
for rate in (50000, 100000, 150000):
    rows.append([
        f"{rate:,.0f} VND/h".replace(",", "."),
        f"{total_hours*rate:,.0f} VND".replace(",", "."),
        f"{contingency*rate:,.0f} VND".replace(",", "."),
        f"{(total_hours+contingency)*rate:,.0f} VND".replace(",", "."),
    ])
add_table(
    ["Hourly rate", "Observed 47.38 h", "Forecast 46.5 h", "Combined"],
    rows,
    [1900, 2400, 2400, 2660], font_size=9.2,
    aligns=[WD_ALIGN_PARAGRAPH.CENTER]*4
)
add_para("These are economic comparison scenarios, not amounts owed to student team members. The SOW must state whether labor is donated, graded coursework, or commercially billable.", italic=True, color=GRAY)

heading("AI/API cost formula", 2)
add_para("For subscription access: AI cash cost = subscription fees + overage fees. For API access: cost = Σ[(uncached input × input rate) + (cached input × cache rate) + (output × output rate)] / 1,000,000. Total tokens alone are insufficient because input, cached input and output rates differ.")

heading("Proposal budget reconciliation", 2)
add_table(
    ["Proposal item", "Budget", "Actual status", "Action"],
    [
        ["AI development tools", "1.60 M VND", "Measured samples used free/trial access", "Insert renewal date, seats and verified post-trial price"],
        ["In-product LLM API", "0.90 M VND", "Not evidenced in current backlog MVP", "Remove or approve as scope change"],
        ["Cloud server + database", "0.80 M VND", "Actual invoices/credits not supplied", "Attach provider, tier, period and invoice"],
        ["Domain + SSL", "0.45 M VND", "Actual purchase not supplied", "Note that managed SSL may be free; record invoice"],
        ["Contingency", "0.50 M VND", "Unspent/unknown", "Define release authority and threshold"],
        ["TOTAL", "4.25 M VND", "Planning baseline, not actual burn", "Maintain planned / committed / actual / forecast-at-completion columns"],
    ],
    [1900, 1300, 2870, 3290], font_size=8.5
)

page_break()
heading("7. Project planning approach", 1)
add_para("A hybrid Kanban-by-batch model is the best fit. User stories flow on a Kanban board, while batch boundaries enforce dependency order and integration gates. This preserves fast AI execution without losing project control.")

heading("Board states and controls", 2)
add_table(
    ["State", "Entry condition", "Exit evidence", "WIP rule"],
    [
        ["Backlog", "Story exists in product backlog", "Product decision resolved", "Unlimited"],
        ["Ready", "AC reviewed; dependency/API contract available", "Owner and test plan assigned", "Max 2 ready per owner"],
        ["In progress", "Owner starts one Codex task for one story", "Code + automated tests locally pass", "Max 1 per owner"],
        ["Review", "PR opened; traceability attached", "Review approved; CI green", "Max 1 per reviewer"],
        ["Integrated", "Merged and migration reproducible", "Cross-module/API test passes", "Batch gate applies"],
        ["Done", "Deployed to agreed environment", "Acceptance criteria verified; no critical/high defect", "Only then count throughput"],
    ],
    [1350, 2870, 3570, 1570], font_size=8.4
)

heading("Batch gates", 2)
add_table(
    ["Gate", "Minimum evidence"],
    [
        ["Batch 1 → 2", "Auth/ownership contract; property-room identifiers; utility/charge schemas; migrations reproducible"],
        ["Batch 2 → 3", "Tenant/lease lifecycle, meter/maintenance APIs and integration fixtures available"],
        ["Batch 3 → 4", "Invoice states, payment/VietQR/reminder behavior, data seeds and audit rules verified"],
        ["Release gate", "Mobile E2E critical paths, deployment smoke test, rollback notes, known-risk acceptance"],
    ],
    [1750, 7610], font_size=9.0
)

heading("Plan cadence", 2)
add_bullet("Daily asynchronous update: Done / next / blocked / telemetry mapping.")
add_bullet("Per-story completion: run tests, map telemetry, attach PR and acceptance evidence.")
add_bullet("Per-batch integration review: 30–45 minutes focused only on API/data/dependency defects.")
add_bullet("Twice-weekly sponsor checkpoint: scope, SOW variance, forecast and decisions; no routine meeting when no decision is needed.")

page_break()
heading("8. Monitoring and control", 1)
heading("Status dashboard", 2)
add_table(
    ["Dimension", "KPI", "Calculation", "Control threshold"],
    [
        ["Scope", "Done stories / committed stories", "Count only global-DoD stories", "Any unapproved story added → scope review"],
        ["Flow", "Cycle time", "In progress → Done", "> 2× rolling median → investigate"],
        ["Quality", "First-pass acceptance", "Stories passing without rework / reviewed", "< 80% → reduce WIP and refine AC"],
        ["Defects", "High/critical open", "Count by severity", "> 0 blocks release"],
        ["AI resource", "Tokens per Done story", "Normalized token fields / Done stories", "> 2× median → inspect context/rework"],
        ["Time", "Forecast variance", "(EAC − baseline) / baseline", "> 15% → reforecast"],
        ["Cost", "Burn variance", "Actual − planned to date", "> 10% or free trial expiry → decision"],
        ["Delivery", "Deployment success", "Successful deployments / attempts", "Any rollback → root-cause note"],
    ],
    [1350, 2200, 3190, 2620], font_size=8.1
)

heading("Demo: collecting evidence for a project status report", 2)
add_number("Create one Codex task for one US and include the US-ID in the task title/prompt.")
add_number("Implement against the story acceptance criteria; run unit/API/integration tests.")
add_number("After completion, run Complete-US.cmd <US-ID> to map turn-level telemetry.")
add_number("Open usage-by-turn.csv / usage-by-conversation.csv / index.html and show tokens and elapsed time.")
add_number("Attach commit, PR, CI result and deployment verification. Only then move the story to Done.")
add_number("Refresh the dashboard: throughput, cycle time, tokens, defects, planned/actual cost and forecast-at-completion.")

callout("Demo evidence path.", r"C:\Users\admin\codex-telemetry\reports\index.html is the readable telemetry view. The raw JSON and CSV files provide auditability; GitHub/CI/deployment evidence proves delivery quality.")

heading("Variance response", 2)
add_bullet("Scope variance: update SOW/change log before implementation.")
add_bullet("Time variance: split oversized stories, reduce WIP, remove blockers, re-estimate remaining work.")
add_bullet("Token variance: start a clean task, reduce context, avoid repeated broad scans, inspect test/rework loops.")
add_bullet("Quality variance: stop starting work; fix acceptance criteria, fixtures and integration tests.")
add_bullet("Cost variance: use approved free tier, cap API spend, or obtain sponsor approval before overage.")

page_break()
heading("9. Statement of Work and oral-defense narrative", 1)
heading("Meaning of the Statement of Work", 2)
add_para("The Statement of Work (SOW) is the agreement baseline between sponsor and development team. It states what will be delivered, what is excluded, who is responsible, what acceptance evidence is required, the schedule/milestones, assumptions, dependencies, cost basis and change-control process. If sponsor expectations and the backlog differ, the team must reconcile them before treating the SOW as approved.")

heading("Minimum SOW fields for RosiHome", 2)
add_table(
    ["Section", "Required decision"],
    [
        ["Scope", "55 backend stories? Mobile only? Which AI product features, if any?"],
        ["Deliverables", "Backend APIs, mobile app, tests, migrations, deployment, documents and demo data"],
        ["Acceptance", "Global DoD + story AC + integration environment + approver"],
        ["Schedule", "Batch milestones, dependency gates and 8–10 week outer deadline"],
        ["Resources", "Five members, AI tools/models, repository, CI, database, hosting"],
        ["Cost", "Cash budget, donated labor, shadow labor and recurring post-trial services"],
        ["Change control", "Who approves new stories, AI features, Web scope or budget variance"],
        ["Operations", "Pilot owner, data/privacy limits, maintenance or shutdown after course"],
    ],
    [1900, 7460], font_size=9.0
)

heading("High-risk scope inconsistencies to resolve", 2)
add_table(
    ["Conflict", "Current evidence", "Required resolution"],
    [
        ["Mobile vs Web", "Backlog PD-07 says Mobile only; proposal timeline says Web + Backend", "Remove Web from SOW or approve it with effort/budget"],
        ["In-product AI", "Proposal budget includes meter-photo AI and weekly AI report; backlog MVP does not define these stories", "Remove 0.9 M API item or create/refine/approve stories"],
        ["Done vs coding complete", "Member reports describe coding batches; global DoD requires deploy and verification", "Publish a story evidence matrix"],
        ["Free tool cost", "Current period is trial/free; proposal budgets paid tools", "Add renewal dates and forecast-at-completion"],
    ],
    [2200, 3570, 3590], font_size=8.6
)

heading("Suggested 90-second oral answer", 2)
callout("Answer structure.", "We began with an 8–10 week sponsor/proposal estimate and a 4.25 million VND tools-and-hosting budget. We then calibrated it using actual AI-assisted backend evidence: 47.38 reported hours and 213.36 million tokens across three work packages. Because token accounting and task types differ, we estimate schedule from Done-story throughput and use tokens only for resource/cost control. Our current three-point backend estimate is 1.42 hours per remaining story before integration; with DoD and risk buffers, 21 nominal stories require about 46.5 agent-supported hours. We manage the work using Kanban inside four dependency batches, and count a story only after tests, review, merge and deployment verification. The SOW is the control baseline that reconciles sponsor scope with backlog scope, especially Mobile versus Web and whether in-product AI is included.")

page_break()
heading("10. Risks, assumptions and actions", 1)
add_table(
    ["ID", "Risk / assumption", "Impact", "Control / owner input"],
    [
        ["R1", "Reported hours use mixed definitions", "Invalid productivity comparison", "Standardize agent, human, blocked and wall time"],
        ["R2", "Token totals use different accounting", "Invalid model/cost comparison", "Capture input/cache/output separately"],
        ["R3", "Story claims may not satisfy global DoD", "Overstated progress", "Attach AC/test/PR/deploy evidence"],
        ["R4", "Dependencies can invalidate ideal parallel time", "Schedule slip", "Use batch gates and contract-first APIs"],
        ["R5", "Free trials expire", "Unexpected recurring cost", "Record expiry, seats, limits and renewal decision"],
        ["R6", "FE and integration evidence missing", "Incomplete project forecast", "Begin identical telemetry/status format for FE"],
        ["R7", "Proposal scope conflicts with backlog", "SOW dispute and budget mismatch", "Sponsor/team scope reconciliation meeting"],
        ["R8", "Infrastructure/service invoices missing", "Actual cost cannot be proven", "Cost ledger with receipt/credit evidence"],
    ],
    [700, 3270, 2160, 3230], font_size=8.4
)

heading("Immediate actions", 2)
add_number("Confirm that the 3 h/5 h/2 h dataset belongs to BE2 Đạt and the 15 h/20 h dataset belongs to BE1 Chí.")
add_number("For each claimed story, mark Coding complete / Tested / Reviewed / Merged / Deployed / Accepted.")
add_number("Reconcile Mobile-only versus Web and in-product AI versus no-AI MVP in the SOW.")
add_number("Record AI trial expiry dates, paid renewal decision and provider billing basis.")
add_number("Adopt the same turn-level telemetry/report format for every member and frontend work.")
add_number("After 5–10 additional Done stories, recompute O/M/P rates and replace the provisional forecast.")

heading("Inputs still required for a defensible final cost", 2)
add_bullet("Human review/prompt-writing time by member and agreed hourly shadow rate.")
add_bullet("Actual hosting, database, storage, email/push, domain, CI and monitoring invoices/credits.")
add_bullet("Post-trial subscription prices, seat counts, renewal dates and API token rate cards.")
add_bullet("Frontend effort, integration effort, blocked time, test failures/rework and deployment history.")
add_bullet("Sponsor-approved SOW and acceptance authority.")

page_break()
heading("Appendix A. Story traceability by backend batch", 1)
story_rows = [
    ["B1", "BE1", "US-AUTH-01→06; US-PROFILE-01", "Registration, login/logout, authorization, password/profile"],
    ["B1", "BE2", "US-PROPERTY-01→02; US-ROOM-01→03", "Property and room setup"],
    ["B1", "BE3", "US-UTILITY-01→02; US-CHARGE-01", "Utility rates and recurring surcharges"],
    ["B2", "BE1", "US-TENANT-01→02; US-LEASE-01→06", "Tenant account and lease lifecycle"],
    ["B2", "BE2", "US-METER-01→03", "Initial/monthly/corrected readings"],
    ["B2", "BE3", "US-MAINT-01→05", "Submit, view, review, status, room history"],
    ["B3", "BE1", "Review, Test, Bug Fix", "Technical work; no story-throughput credit by itself"],
    ["B3", "BE2", "US-INVOICE-01→04", "Generate, view, download, review/send invoices"],
    ["B3", "BE3", "US-VIETQR-01→02; US-PAYMENT-01→03; US-REMINDER-01→02", "Payment configuration, proof, verification/history, reminders"],
    ["B4", "BE1", "US-DASH-01→02", "Occupancy and revenue dashboard"],
    ["B4", "BE2", "US-DASH-03→04", "Debt and lease-expiry dashboard"],
    ["B4", "BE3", "US-REPORT-01→05", "Reporting period, finance, occupancy, maintenance, PDF"],
]
add_table(["Batch", "Owner", "Stories / task", "Product outcome"], story_rows, [800, 900, 3730, 3930], font_size=8.5)

heading("Appendix B. Calculation ledger", 1)
add_table(
    ["Calculation", "Result"],
    [
        ["Minh total", "56.564626 M tokens; 8,569.2 s = 2.3803 h"],
        ["Đạt total", "6.3 + 13.0 + 2.5 = 21.8 M tokens; 3 + 5 + 2 = 10 h"],
        ["Chí total", "135 M tokens; 15 + 20 = 35 h"],
        ["Observed aggregate", "213.364626 M tokens; 47.3803 h"],
        ["Provisional covered stories", "5 + 14 + (3 named Auth + 12 Batch-2) = 34"],
        ["Remaining nominal scope", "55 − 34 = 21 stories"],
        ["PERT expected rate", "(0.65 + 4×1.39 + 2.33) / 6 = 1.42 h/US"],
        ["Planning commitment", "21 × 1.42 × 1.30 × 1.20 = 46.5 h"],
    ],
    [3200, 6160], font_size=9.0
)

heading("Appendix C. Source register", 1)
add_table(
    ["Source", "Use"],
    [
        [r"D:\Downloads\Phân công.txt", "Governing BE/FE batch allocation"],
        [r"docs\product_backlog.md", "Story outcomes, dependencies, product decisions and global DoD"],
        [r"docs\proposal\time.md", "8–10 week proposal schedule baseline"],
        [r"docs\proposal\cost.md and budget.md", "4.25 M VND budget baseline"],
        [r"C:\Users\admin\codex-telemetry\reports\*", "Minh turn-level token/time evidence"],
        ["Member declarations supplied in conversation", "Đạt/Chí aggregate time and token evidence"],
    ],
    [3900, 5460], font_size=8.8
)

# Avoid splitting table rows and set paragraph widow/orphan controls.
for table in doc.tables:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
for p in doc.paragraphs:
    p_pr = p._p.get_or_add_pPr()
    widow = OxmlElement("w:widowControl")
    widow.set(qn("w:val"), "1")
    p_pr.append(widow)

doc.core_properties.title = "RosiHome Cost, Time & Resources Report"
doc.core_properties.subject = "Software Project Estimation, Planning, Monitoring and Control"
doc.core_properties.author = "RosiHome Team"
doc.core_properties.keywords = "RosiHome, estimation, AI coding, cost, time, resources, monitoring"
doc.save(OUT)
print(OUT)
